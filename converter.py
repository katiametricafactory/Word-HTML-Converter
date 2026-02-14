import re
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from html import escape
import tempfile
import subprocess
import os

FONT_DIV = 'font-family: Arial; font-size: 11pt;'

def is_number_cell(text):
    return bool(re.search(r'\d+,\d+|\d+\.\d+|€', text))

def convert_doc_to_docx(input_path):
    output_dir = os.path.dirname(input_path)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", input_path, "--outdir", output_dir],
        check=True
    )
    return input_path.replace(".doc", ".docx")

def build_inline_table(table):
    rows = table.rows
    col_count = len(rows[0].cells)

    html = []
    html.append('<br>')
    html.append('<table style="border-collapse: collapse; width: 100%; border: 1px solid #000; font-size: 9pt; table-layout: auto;" border="1">')

    html.append('<thead>')
    html.append(f'''
<tr>
<th colspan="{col_count}" 
style="border:1px solid #000; background:#0a8f08; color:#fff; 
text-align:center; padding:4px; font-size:9pt;">
{escape(rows[0].cells[0].text)}
</th>
</tr>
''')

    html.append('<tr>')
    for cell in rows[1].cells:
        html.append(f'''
<th style="border:1px solid #000; background:#c6efce; 
text-align:left; padding:4px; font-size:9pt; white-space: nowrap;">
{escape(cell.text)}
</th>
''')
    html.append('</tr></thead><tbody>')

    for row in rows[2:]:
        texts = [c.text.strip() for c in row.cells]
        is_total = any("total" in t.lower() or "totais" in t.lower() for t in texts)

        html.append('<tr>')
        for cell in row.cells:
            txt = cell.text.strip()
            is_number = is_number_cell(txt)

            align = "right" if is_number else "left"
            nowrap = "white-space: nowrap;" if is_number else ""
            bg = "background:#c9c9c9;" if is_total else ""

            html.append(f'''
<td style="border:1px solid #000; padding:3px; 
text-align:{align}; font-size:9pt; {nowrap} {bg}">
{escape(txt) if txt else '&nbsp;'}
</td>
''')
        html.append('</tr>')

    html.append('</tbody></table><br><br>')
    return "\n".join(html)

def convert_word_to_html(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name

    if tmp_path.lower().endswith(".doc"):
        tmp_path = convert_doc_to_docx(tmp_path)

    doc = Document(tmp_path)

    html_parts = []
    html_parts.append(f'<div style="{FONT_DIV}">')

    table_index = 0

    for child in doc.element.body.iterchildren():

        if child.tag == qn('w:p'):
            para = Paragraph(child, doc)
            text = para.text.strip()

            if not text:
                html_parts.append("<p>&nbsp;</p>")
            else:
                html_parts.append(f"<p>{escape(text)}</p>")

        elif child.tag == qn('w:tbl'):
            if table_index < len(doc.tables):
                html_parts.append(build_inline_table(doc.tables[table_index]))
                table_index += 1

    html_parts.append("</div>")

    return "\n".join(html_parts)
